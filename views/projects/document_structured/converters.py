import os
import tempfile
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from io import BytesIO
from xhtml2pdf import pisa

class DocumentConverter:
    @staticmethod
    def markdown_to_pdf(markdown_text, output_path):
        """Convert markdown text to PDF using xhtml2pdf"""
        # First convert markdown to HTML
        html = markdown.markdown(markdown_text)
        
        # Create PDF
        with open(output_path, "w+b") as output_file:
            pisa_status = pisa.CreatePDF(html, dest=output_file)
        
        return pisa_status.err

    @staticmethod
    def markdown_to_docx(markdown_text, output_path):
        """Convert markdown text to DOCX"""
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading('Structured Document', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process markdown content
        lines = markdown_text.split('\n')
        
        for line in lines:
            if line.startswith('# '):
                # H1 heading
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # H2 heading
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # H3 heading
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # List item
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif line.strip() == '':
                # Empty line
                doc.add_paragraph()
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
        
        doc.save(output_path)

    @staticmethod
    def create_download_file(markdown_text, format_type):
        """Create a temporary file for download"""
        if format_type == "pdf":
            ext = "pdf"
            converter = DocumentConverter.markdown_to_pdf
        elif format_type == "docx":
            ext = "docx"
            converter = DocumentConverter.markdown_to_docx
        else:
            raise ValueError("Unsupported format type")
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}')
        temp_path = temp_file.name
        temp_file.close()
        
        # Convert and save
        converter(markdown_text, temp_path)
        
        return temp_path